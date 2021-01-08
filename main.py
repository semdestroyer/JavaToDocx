
from pathlib import Path
from docx import Document
from tkinter import *
import os
from tkinter import filedialog
class main:
    def __init__(self):
        self.inpt = ''
        self.outp = ''


    def openjava(self):
        inp = filedialog.askdirectory()
        self.inpt = inp
        print(self.inpt)

    def savedocx(self):
        files = [('docx', '*.docx*')]
        out = filedialog.asksaveasfile(filetypes = files, defaultextension = '.docx')
        self.outp = out.name
        lbl = Label(text=self.outp)
        lbl.pack()


    def convert(self):
        document = Document()
        fils = os.listdir(self.inpt)

        pathlist = Path(self.inpt).glob('**/*.java')
        for p in pathlist:
            filstr = open(p,encoding='utf-8')
            document.add_heading(str(p), 0)
            document.add_paragraph(filstr)
            document.add_page_break()




        document.save(str(self.outp))



    def execut(self):
        window = Tk()
        window.title("Java в docx")

        butt = Button(window,text="Открыть папку с java файлами", command = self.openjava)
        butt2 = Button(window,text="выбрать путь для выходного docx", command = self.savedocx)

        butt3 = Button(window,text="конвертировать", command = self.convert)
        butt.pack()
        butt2.pack()
        butt3.pack()

        window.mainloop()


m = main()
m.execut()
